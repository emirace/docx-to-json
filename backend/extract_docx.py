from docx import Document
import json
import sys

def extract_docx_to_json(docx_path):
    doc = Document(docx_path)
    data = []

    for paragraph in doc.paragraphs:
        style = paragraph.style.name
        runs = [
            {
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'color': run.font.color.rgb if run.font.color else None
            }
            for run in paragraph.runs
        ]
        data.append({'type': 'paragraph', 'style': style, 'content': runs})

    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [{'text': cell.text, 'style': cell.paragraphs[0].style.name} for cell in row.cells]
            table_data.append(row_data)
        data.append({'type': 'table', 'content': table_data})

    return json.dumps(data)

if __name__ == "__main__":
    docx_path = sys.argv[1]
    json_data = extract_docx_to_json(docx_path)
    print(json_data)
